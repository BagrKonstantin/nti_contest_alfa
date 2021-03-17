from stud_card import Ui_MainWindow as Dialog_ui
from decan_vision import Ui_MainWindow
from work_with_app import Ui_MainWindow as Ui_MainWindow2
from PyQt5.QtWidgets import QMainWindow, QMessageBox
from PyQt5 import QtWidgets, QtGui
from PyQt5.QtGui import QPixmap
import time
from email.mime.text import MIMEText
from PyQt5.QtWidgets import *
import smtplib
import sqlite3
import smtplib
import os
import docx



class UI_Task1(QMainWindow, Ui_MainWindow):
    def __init__(self, path):
        super(UI_Task1, self).__init__()
        self.setupUi(self)
        self.setWindowTitle('enter data')
        self.path = path

        self.tableWidget_5.cellClicked.connect(self.open_tab)
        self.comboBox_2.addItems(['', 'ПМ', 'ИВТ', 'АУТС', 'ИкТ', 'ЗСС'])

        self.pushButton_2.clicked.connect(self.update_table2)
        self.tableWidget_5.horizontalHeader().setSectionResizeMode(0, QtWidgets.QHeaderView.Stretch)
        self.update_data()



    def update_table2(self):
        try:
            con = sqlite3.connect(self.path)
            cur = con.cursor()
            data = cur.execute("""Select * from user where napravlenie = '{}'""".format(self.comboBox_2.currentText())).fetchall()
            con.close()
            self.tableWidget_5.setRowCount(len(data))
            for i in range(len(data)):
                self.tableWidget_5.setItem(i, 0, QTableWidgetItem())

                # print(sum(data[i][23] + data[i][24] + data[i][25] + data[i][26]))

                self.tableWidget_5.item(i, 0).setText(str(data[i][1] + " " + data[i][2] + " " + data[i][3]))

        except Exception as err:
            print(err)



    def update_data(self):
        con = sqlite3.connect(self.path)
        cur = con.cursor()
        self.data = cur.execute("""Select * from user""").fetchall()
        con.close()

    def open_tab(self):
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Question)
        # msg.setIconPixmap(pixmap)  # Своя картинка

        msg.setWindowTitle("Абитуриент")
        msg.setText("Вы хотите посмотреть данные абитуриента?")
        # msg.setInformativeText("Хотите сохранить информацию перед выходом?")

        msg.addButton('Отмена', QMessageBox.YesRole)
        ok_button = msg.addButton('Да', QMessageBox.AcceptRole)
        # close_button = msg.addButton('Закрыть', QMessageBox.RejectRole)

        msg.exec()
        if msg.clickedButton() == ok_button:
            print("opendialog")
            try:
                print(self.tableWidget_5.verticalHeader().sortIndicatorSection())
                dialog = Dialog(self, self.data[self.tableWidget_5.verticalHeader().sortIndicatorSection() - 1])
                dialog.show()
                print("no error")
            except Exception as error:
                print(error)


class Dialog(QMainWindow, Dialog_ui):
    def __init__(self, parent=None, *args):
        parent.setDisabled(True)
        self.parent = parent
        super(Dialog, self).__init__(parent)
        self.setupUi(self)
        self.setWindowTitle('Dialog')

        # doc = docx.Document()
        #
        # # добавляем первый параграф
        # doc.add_paragraph('Здравствуй, мир!')
        #
        # # добавляем еще два параграфа
        # par1 = doc.add_paragraph('Это второй абзац.')
        # par2 = doc.add_paragraph('Это третий абзац.')
        #
        # # добавляем текст во второй параграф
        # par1.add_run(' Этот текст был добавлен во второй абзац.')
        #
        # # добавляем текст в третий параграф
        # par2.add_run(' Добавляем текст в третий абзац.').bold = True
        #
        # doc.save('helloworld.docx')
        # os.startfile('helloworld.docx')




        self.data = args[0]

        self.lineEdit_name_2.setText(str(self.data[1]))
        self.lineEdit_secondname_2.setText(str(self.data[3]))
        self.lineEdit_surname_2.setText(str(self.data[2]))
        self.lineEdit_sex_2.setText("Мужской" if self.data[4] == 0 else "Женский")
        self.lineEdit_email_2.setText(str(self.data[5]))
        # password
        self.lineEdit_birht_date_2.setText(str(self.data[7]))
        self.lineEdit_phone_2.setText(str(self.data[8]))
        # self.lineEdit_place_birth_2.setText(str(self.data[9]))
        # self.radioButton_hostel_2.setChecked(True if self.data[10] else False)
        # фото
        self.lineEdit_series_2.setText(str(self.data[12]))
        self.lineEdit_pass_number_2.setText(str(self.data[13]))
        self.lineEdit_given_by_2.setText(str(self.data[14]))
        self.lineEdit_code_2.setText(str(self.data[15]))
        self.lineEdit_given_date_2.setText(str(self.data[16]))
        # фото паспорт
        self.lineEdit_adress_2.setText(str(self.data[18]))
        # фото согласие
        self.lineEdit_number_of_doc_frame_3.setText(str(self.data[20]))
        # фото аттестат
        # self.lineEdit_teach_form.setText("бюджет" if self.data[22] else "платное")
        # self.lineEdit_rus_2.setText(str(self.data[23]))
        # self.lineEdit_math_2.setText(str(self.data[24]))
        # self.lineEdit_inf.setText(str(self.data[25]))
        # self.lineEdit_fizika.setText(str(self.data[26]))
        # self.lineEdit_achiev.setText(str(self.data[27]))
        # self.lineEdit_index_2.setText(str(self.data[34]))
        self.lineEdit_direction.setText(str(self.data[33]))
        # фото достижение

        self.user_photo_path_1 = 'user_photos/' + self.data[11]
        self.pixmap = QPixmap(self.user_photo_path_1)
        self.user_photo.setPixmap(self.pixmap)


        self.first_page_path_2 = 'passports/' + self.data[17]
        self.pixmap = QPixmap(self.first_page_path_2)
        self.first_page.setPixmap(self.pixmap)





        self.pushButton.clicked.connect(self.back)

        self.pushButton_user.clicked.connect(self.open0)
        self.pushButton_first_page.clicked.connect(self.open1)

        self.data_file = ['user_photos/' + self.data[11], 'passports/' + self.data[17], 'agreements/' + self.data[19], 'attestats/' + self.data[21], 'achiev/' + self.data[28]]

    def open0(self):
        print(os.getcwd())
        try:
            os.startfile(os.getcwd() + "/" + self.data_file[0])
        except Exception as err:
            print(err)
    def open1(self):
        os.startfile(os.getcwd() + "/" + self.data_file[1])


    def closeEvent(self, event):
        self.parent.setDisabled(False)


    def back(self):
        self.close()
        self.parent.setDisabled(False)


if __name__ == "__main__":
    import sys

    app = QtWidgets.QApplication(sys.argv)
    mainWindow = UI_Task1("bd.db")
    mainWindow.show()
    sys.exit(app.exec())
